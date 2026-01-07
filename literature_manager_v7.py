import re
import json
import os
import webbrowser
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
from docx import Document


# ===========================================================
#                 改进版作者解析
# ===========================================================
def smart_parse_authors(raw: str):
    if not raw:
        return []
    s = raw.replace(" and ", " ; ").replace(" & ", " ; ").replace("；", ";")

    parts = [p.strip(" ,.") for p in s.split(";") if p.strip(" ,.")]
    authors = []
    for p in parts:
        if "," in p:      # Reeves, C. R.
            fam, giv = [x.strip() for x in p.split(",", 1)]
        else:             # C. R. Reeves
            toks = p.split()
            fam = toks[-1]
            giv = " ".join(toks[:-1])
        authors.append({"family": fam, "given": giv})
    return authors


# ===========================================================
#                 文献自动识别（APA/MLA/GB/T）
# ===========================================================
def parse_citation(raw: str):
    text = raw.replace("\n", " ").strip()
    result = {
        "authors": [],
        "title": "",
        "journal": "",
        "volume": "",
        "issue": "",
        "pages": "",
        "year": "",
        "doi": "",
        "url": "",
        "raw_text": raw.strip(),
        "files": []
    }

    # DOI
    doi = re.search(r"(10\.\d{4,9}/[-._;()/:A-Za-z0-9]+)", text)
    if doi:
        result["doi"] = doi.group(1)

    # URL
    url = re.search(r"(https?://[^\s]+)", text)
    if url:
        result["url"] = url.group(1)

    # 页码
    pages = re.search(r"(\d{1,5}\s*[-–]\s*\d{1,5})", text)
    if pages:
        result["pages"] = pages.group(1).replace(" ", "")

    # APA: Authors (Year). Title. Journal
    apa = re.compile(
        r"^(?P<authors>.+?)\s*\((?P<year>19\d{2}|20\d{2})\)\.\s*(?P<title>.+?)\.\s*(?P<rest>.+)$"
    )
    m = apa.match(text)
    if m:
        result["authors"] = smart_parse_authors(m.group("authors"))
        result["year"] = m.group("year")
        result["title"] = m.group("title").strip()
        result["journal"] = m.group("rest").strip()
        return result

    # 句号分段
    seg = [s.strip() for s in text.split(".") if s.strip()]
    if len(seg) >= 2:
        result["authors"] = smart_parse_authors(seg[0])
        result["title"] = seg[1]
        if len(seg) >= 3:
            result["journal"] = seg[2]
        y = re.search(r"(19\d{2}|20\d{2})", text)
        if y:
            result["year"] = y.group(1)
        return result

    # 逗号回退
    parts = [p.strip() for p in text.split(",") if p.strip()]
    if parts:
        result["authors"] = smart_parse_authors(parts[0])
    if len(parts) >= 2:
        result["title"] = parts[1]
    if len(parts) >= 3:
        result["journal"] = parts[2]

    year = re.search(r"(19\d{2}|20\d{2})", text)
    if year:
        result["year"] = year.group(1)

    return result


# ===========================================================
#                        DOCX 工具
# ===========================================================
def apa_citation(entry):
    if not entry["authors"]:
        return f"(Unknown, {entry['year']})"
    if len(entry["authors"]) == 1:
        return f"({entry['authors'][0]['family']}, {entry['year']})"
    if len(entry["authors"]) == 2:
        return f"({entry['authors'][0]['family']} & {entry['authors'][1]['family']}, {entry['year']})"
    return f"({entry['authors'][0]['family']} et al., {entry['year']})"


def replace_ids(doc, library, style):
    id_map = {}
    counter = 1

    for p in doc.paragraphs:
        ids = re.findall(r"\[id:(.*?)\]", p.text)
        for rid in ids:
            if rid not in id_map:
                id_map[rid] = counter
                counter += 1

    for p in doc.paragraphs:
        for rid, num in id_map.items():
            entry = next((x for x in library if x["id"] == rid), None)
            if not entry:
                continue
            rep = apa_citation(entry) if style == "APA7" else f"[{num}]"
            p.text = p.text.replace(f"[id:{rid}]", rep)

    return id_map


def insert_refs(doc, library, id_map, style):
    doc.add_heading("References" if style == "APA7" else "参考文献", level=1)
    for rid, num in sorted(id_map.items(), key=lambda x: x[1]):
        e = next((x for x in library if x["id"] == rid), None)
        if not e:
            continue
        line = f"[{num}] {e['title']}. {e['journal']} ({e['year']})."
        doc.add_paragraph(line)


# ===========================================================
#                    GUI 主程序类
# ===========================================================
class LiteratureManager:
    def __init__(self):
        self.window = tk.Tk()
        self.window.title("文献管理器@AskSpace.cn（自动识别 + 原文保存 + 多附件 + 自动保存 TXT）")
        self.window.geometry("1200x750")

        self.library = []
        self.current_style = "APA7"
        self.next_id = 1
        self.current_txt_path = None

        self.build_menubar()
        self.build_layout()

        self.window.mainloop()

    # ------------------------ 菜单栏 ------------------------
    def build_menubar(self):
        menubar = tk.Menu(self.window)

        # 文件
        m_file = tk.Menu(menubar, tearoff=0)
        m_file.add_command(label="导入文献库 TXT", command=self.load_txt)
        m_file.add_command(label="另存为文献库 TXT", command=self.save_as_txt)
        menubar.add_cascade(label="文件", menu=m_file)

        # 文献库
        m_lib = tk.Menu(menubar, tearoff=0)
        m_lib.add_command(label="新建空文献", command=self.add_entry)
        m_lib.add_command(label="从字符串自动识别", command=self.add_from_string)
        m_lib.add_command(label="删除文献", command=self.delete_entry)
        menubar.add_cascade(label="文献库", menu=m_lib)

        # 引用格式
        m_style = tk.Menu(menubar, tearoff=0)
        for s in ["APA7", "GB/T", "IEEE"]:
            m_style.add_command(label=s, command=lambda x=s: self.set_style(x))
        menubar.add_cascade(label="引用格式", menu=m_style)

        # DOCX
        m_docx = tk.Menu(menubar, tearoff=0)
        m_docx.add_command(label="处理 DOCX 引用", command=self.process_docx)
        menubar.add_cascade(label="DOCX 工具", menu=m_docx)

        # 帮助
        m_help = tk.Menu(menubar, tearoff=0)
        m_help.add_command(label="AskSpace版权所有,联系我们SpaceAero@163.com", command=lambda: webbrowser.open("http://askspace.cn/wx.html"))
        m_help.add_command(label="AskSpace.cn 帮助中心", command=lambda: webbrowser.open("http://askspace.cn/wx.html"))
        menubar.add_cascade(label="帮助", menu=m_help)

        self.window.config(menu=menubar)

    # ------------------------ 主布局 ------------------------
    def build_layout(self):
        main = tk.Frame(self.window)
        main.pack(fill="both", expand=True)

        # 左侧视图
        left = tk.Frame(main)
        left.pack(side="left", fill="both")

        tk.Label(left, text="文献库列表：", font=("Arial", 12, "bold")).pack(anchor="w")

        btns = tk.Frame(left)
        btns.pack(anchor="w", pady=3)
        tk.Button(btns, text="新建", command=self.add_entry).pack(side="left", padx=2)
        tk.Button(btns, text="自动识别添加", command=self.add_from_string).pack(side="left", padx=2)
        tk.Button(btns, text="删除", command=self.delete_entry).pack(side="left", padx=2)

        columns = ("id", "authors", "title", "year", "doi")
        self.tree = ttk.Treeview(left, columns=columns, show="headings", height=30)

        for col, w in zip(columns, [60, 220, 300, 60, 220]):
            self.tree.heading(col, text=col)
            self.tree.column(col, width=w)

        self.tree.pack(fill="both")
        self.tree.bind("<<TreeviewSelect>>", self.on_select)

        # 右侧编辑区
        right = tk.Frame(main)
        right.pack(side="left", fill="both", expand=True)

        tk.Label(right, text="文献信息编辑", font=("Arial", 12, "bold")).pack(anchor="w")

        self.fields = {}
        labels = [
            ("标题", "title"),
            ("作者", "authors"),
            ("年份", "year"),
            ("期刊", "journal"),
            ("卷", "volume"),
            ("期", "issue"),
            ("页码", "pages"),
            ("DOI", "doi"),
            ("URL", "url")
        ]

        for lab, key in labels:
            row = tk.Frame(right)
            row.pack(fill="x", pady=3)
            tk.Label(row, text=f"{lab}：", width=10).pack(side="left")
            ent = tk.Entry(row)
            ent.pack(side="left", fill="x", expand=True)
            self.fields[lab] = ent

        # 原文
        tk.Label(right, text="原始引用：", font=("Arial", 10, "bold")).pack(anchor="w")
        self.raw_box = scrolledtext.ScrolledText(right, height=5)
        self.raw_box.pack(fill="x")

        # 附件部分
        attach = tk.Frame(right)
        attach.pack(anchor="w", pady=5)
        tk.Button(attach, text="关联此引用的附件", command=self.add_file).pack(side="left", padx=3)
        tk.Button(attach, text="查看附件", command=self.show_files).pack(side="left", padx=3)

        # 保存按钮
        tk.Button(right, text="保存修改", command=self.save_changes).pack(anchor="w", pady=10)

        # 日志
        tk.Label(self.window, text="日志：", font=("Arial", 11, "bold")).pack(anchor="w")
        self.log = scrolledtext.ScrolledText(self.window, height=7)
        self.log.pack(fill="both")

    # ======================================================
    #                 文献操作函数
    # ======================================================
    def logmsg(self, t):
        self.log.insert(tk.END, t + "\n")
        self.log.see(tk.END)

    def refresh_tree(self):
        self.tree.delete(*self.tree.get_children())
        for e in self.library:
            au = "; ".join([f"{x['family']} {x['given']}" for x in e["authors"]])
            self.tree.insert("", "end", values=(e["id"], au, e["title"], e["year"], e["doi"]))

    def add_entry(self):
        e = {
            "id": f"ref{self.next_id:03d}",
            "authors": [],
            "title": "",
            "journal": "",
            "volume": "",
            "issue": "",
            "pages": "",
            "year": "",
            "doi": "",
            "url": "",
            "raw_text": "",
            "files": []
        }
        self.next_id += 1
        self.library.append(e)
        self.refresh_tree()
        self.auto_save()
        self.logmsg(f"新建文献：{e['id']}")

    def delete_entry(self):
        sel = self.tree.selection()
        if not sel:
            return
        index = self.tree.index(sel[0])
        rid = self.library[index]["id"]
        self.library.pop(index)
        self.refresh_tree()
        self.auto_save()
        self.logmsg(f"已删除：{rid}")

    def add_from_string(self):
        raw = simpledialog.askstring("粘贴文献", "请输入参考文献引用：")
        if not raw:
            return

        e = parse_citation(raw)
        e["id"] = f"ref{self.next_id:03d}"
        self.next_id += 1

        self.library.append(e)
        self.refresh_tree()
        self.auto_save()
        self.logmsg(f"自动识别添加：{e['id']}")

    # ======================================================
    #                    编辑行为
    # ======================================================
    def on_select(self, evt):
        sel = self.tree.selection()
        if not sel:
            return

        index = self.tree.index(sel[0])
        e = self.library[index]

        # 简短字段
        def setf(lab, key):
            self.fields[lab].delete(0, tk.END)
            self.fields[lab].insert(0, e.get(key, ""))

        setf("标题", "title")

        # 作者 family,given;family,given
        if e["authors"]:
            s = ";".join([f"{x['family']},{x['given']}" for x in e["authors"]])
        else:
            s = ""
        self.fields["作者"].delete(0, tk.END)
        self.fields["作者"].insert(0, s)

        setf("年份", "year")
        setf("期刊", "journal")
        setf("卷", "volume")
        setf("期", "issue")
        setf("页码", "pages")
        setf("DOI", "doi")
        setf("URL", "url")

        # 原文
        self.raw_box.delete(1.0, tk.END)
        self.raw_box.insert(tk.END, e["raw_text"])

    # ======================================================
    #                保存修改（含自动保存 TXT）
    # ======================================================
    def save_changes(self):
        sel = self.tree.selection()
        if not sel:
            return

        index = self.tree.index(sel[0])
        e = self.library[index]

        e["title"] = self.fields["标题"].get().strip()
        e["journal"] = self.fields["期刊"].get().strip()
        e["volume"] = self.fields["卷"].get().strip()
        e["issue"] = self.fields["期"].get().strip()
        e["pages"] = self.fields["页码"].get().strip()
        e["year"] = self.fields["年份"].get().strip()
        e["doi"] = self.fields["DOI"].get().strip()
        e["url"] = self.fields["URL"].get().strip()
        e["raw_text"] = self.raw_box.get(1.0, tk.END).strip()

        # 作者
        raw = self.fields["作者"].get().strip()
        authors = []
        if raw:
            parts = [x.strip() for x in raw.split(";") if x.strip()]
            for p in parts:
                if "," in p:
                    fam, giv = [x.strip() for x in p.split(",", 1)]
                else:
                    toks = p.split()
                    fam = toks[-1]
                    giv = " ".join(toks[:-1])
                authors.append({"family": fam, "given": giv})
        e["authors"] = authors

        self.refresh_tree()
        self.auto_save()
        self.logmsg(f"已保存修改：{e['id']}")

    # ======================================================
    #                    附件（多文件）
    # ======================================================
    def add_file(self):
        sel = self.tree.selection()
        if not sel:
            return

        index = self.tree.index(sel[0])
        e = self.library[index]

        paths = filedialog.askopenfilenames()
        if not paths:
            return

        e["files"].extend(list(paths))
        self.auto_save()
        self.logmsg(f"已添加附件（{len(paths)} 个）")

    def open_file_direct(self, path):
        if not os.path.exists(path):
            messagebox.showerror("错误", f"文件不存在：\n{path}")
            return

        try:
            if os.name == "nt":
                os.startfile(path)
            elif os.uname().sysname == "Darwin":
                os.system(f"open '{path}'")
            else:
                os.system(f"xdg-open '{path}'")
        except Exception as exc:
            messagebox.showerror("无法打开", f"{exc}")

    def show_files(self):
        sel = self.tree.selection()
        if not sel:
            return
        index = self.tree.index(sel[0])
        e = self.library[index]

        if not e["files"]:
            messagebox.showinfo("附件", "无附件")
            return

        win = tk.Toplevel(self.window)
        win.title(f"{e['id']} 附件列表")
        win.geometry("600x300")

        tk.Label(win, text=f"文献 {e['id']} 的附件：", font=("Arial", 11, "bold")).pack(anchor="w", pady=5)

        frame = tk.Frame(win)
        frame.pack(fill="both", expand=True)

        canvas = tk.Canvas(frame)
        canvas.pack(side="left", fill="both", expand=True)

        sb = tk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        sb.pack(side="right", fill="y")
        canvas.configure(yscrollcommand=sb.set)

        inner = tk.Frame(canvas)
        canvas.create_window((0, 0), window=inner, anchor="nw")

        inner.bind(
            "<Configure>",
            lambda evt: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        # 添加行
        for path in e["files"]:
            row = tk.Frame(inner)
            row.pack(fill="x", pady=4)

            tk.Label(row, text=os.path.basename(path), anchor="w").pack(side="left", padx=5)

            tk.Button(
                row, text="打开",
                command=lambda p=path: self.open_file_direct(p)
            ).pack(side="right", padx=5)

    # ======================================================
    #                      TXT 保存
    # ======================================================
    def auto_save(self):
        if self.current_txt_path is None:
            path = filedialog.asksaveasfilename(defaultextension=".txt")
            if not path:
                self.logmsg("未选择 文献库 TXT 保存路径")
                return
            self.current_txt_path = path

        with open(self.current_txt_path, "w", encoding="utf8") as f:
            for e in self.library:
                f.write(json.dumps(e, ensure_ascii=False) + "\n")

        self.logmsg(f"自动保存：{self.current_txt_path}")

    def save_as_txt(self):
        path = filedialog.asksaveasfilename(defaultextension=".txt")
        if not path:
            return
        self.current_txt_path = path
        self.auto_save()

    def load_txt(self):
        path = filedialog.askopenfilename(filetypes=[("TXT", "*.txt")])
        if not path:
            return
        with open(path, "r", encoding="utf8") as f:
            self.library = [json.loads(x) for x in f]

        # 修复 next_id
        self.next_id = max([int(e["id"][3:]) for e in self.library] + [0]) + 1
        self.current_txt_path = path

        self.refresh_tree()
        self.logmsg("加载文献库 TXT 完成")

    # ======================================================
    #                       DOCX
    # ======================================================
    def process_docx(self):
        path = filedialog.askopenfilename(filetypes=[("DOCX", "*.docx")])
        if not path:
            return

        doc = Document(path)
        id_map = replace_ids(doc, self.library, self.current_style)
        insert_refs(doc, self.library, id_map, self.current_style)

        out = path.replace(".docx", f"_output_{self.current_style}.docx")
        doc.save(out)
        self.logmsg(f"已生成：{out}")

    # ======================================================
    #                     引用格式
    # ======================================================
    def set_style(self, style):
        self.current_style = style
        self.logmsg(f"引用格式切换为：{style}")


# ===========================================================
# 主入口
# ===========================================================
if __name__ == "__main__":
    LiteratureManager()
